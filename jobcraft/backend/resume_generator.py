"""
PDF resume generator — takes the structured tailored-resume JSON
and produces a clean, professional PDF using ReportLab.
"""

import logging
import re
import uuid
from pathlib import Path
from typing import Dict

from docx import Document
from docx.shared import Pt
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, HRFlowable
)
from reportlab.lib.enums import TA_CENTER

from config import RESUMES_TAILORED_DIR

logger = logging.getLogger(__name__)


def _safe_stem(company: str) -> str:
    """Create a safe filename stem from company name."""
    clean = (company or "company").strip().replace(" ", "_")
    return "".join(ch for ch in clean if ch.isalnum() or ch in {"_", "-"})[:60] or "company"


def _clean_text(value: str) -> str:
    """Normalize whitespace and strip control characters for Word/PDF safety."""
    s = (value or "").replace("\r", " ").replace("\n", " ")
    s = "".join(ch for ch in s if ch.isprintable())
    return " ".join(s.split()).strip()


def _add_paragraph(doc: Document, text: str, *, bold: bool = False, size_pt: int = 11, style: str | None = None):
    """Add a paragraph with explicit run formatting (prevents invisible text)."""
    text = _clean_text(text)
    if not text:
        return None
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bool(bold)
    run.font.size = Pt(size_pt)
    return p


def _build_styles():
    """Custom paragraph styles for the resume PDF."""
    styles = getSampleStyleSheet()

    styles.add(ParagraphStyle(
        "ResumeName",
        parent=styles["Heading1"],
        fontSize=18,
        spaceAfter=4,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#1a1a2e"),
    ))
    styles.add(ParagraphStyle(
        "ResumeSummary",
        parent=styles["Normal"],
        fontSize=10,
        leading=14,
        spaceAfter=10,
        textColor=colors.HexColor("#333333"),
    ))
    styles.add(ParagraphStyle(
        "SectionTitle",
        parent=styles["Heading2"],
        fontSize=12,
        spaceBefore=12,
        spaceAfter=4,
        textColor=colors.HexColor("#6366F1"),
    ))
    styles.add(ParagraphStyle(
        "CompanyRole",
        parent=styles["Normal"],
        fontSize=11,
        leading=14,
        spaceBefore=6,
        textColor=colors.HexColor("#1a1a2e"),
        fontName="Helvetica-Bold",
    ))
    styles.add(ParagraphStyle(
        "ResumeBullet",
        parent=styles["Normal"],
        fontSize=10,
        leading=13,
        leftIndent=18,
        bulletIndent=6,
        textColor=colors.HexColor("#333333"),
    ))
    styles.add(ParagraphStyle(
        "SkillText",
        parent=styles["Normal"],
        fontSize=10,
        leading=13,
        textColor=colors.HexColor("#333333"),
    ))
    return styles


def generate_pdf(tailored_resume: Dict, job_title: str = "", company: str = "") -> Path:
    """
    Create a professional PDF from the tailored resume JSON.
    Returns the Path to the generated file.
    """
    RESUMES_TAILORED_DIR.mkdir(parents=True, exist_ok=True)
    filename = f"{uuid.uuid4().hex[:12]}_{_safe_stem(company)}.pdf"
    output_path = RESUMES_TAILORED_DIR / filename

    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        topMargin=0.6 * inch,
        bottomMargin=0.6 * inch,
        leftMargin=0.7 * inch,
        rightMargin=0.7 * inch,
    )

    styles = _build_styles()
    story = []

    # --- Summary ---
    summary = tailored_resume.get("summary", "")
    if summary:
        story.append(Paragraph(summary, styles["ResumeSummary"]))
        story.append(Spacer(1, 6))

    # --- Horizontal rule ---
    story.append(HRFlowable(
        width="100%", thickness=0.5,
        color=colors.HexColor("#d1d5db"), spaceAfter=8
    ))

    # --- Experience ---
    experience = tailored_resume.get("experience", [])
    if experience:
        story.append(Paragraph("EXPERIENCE", styles["SectionTitle"]))
        for exp in experience:
            role_line = f"{exp.get('role', '')} — {exp.get('company', '')}"
            dates = exp.get("dates", "")
            story.append(Paragraph(role_line, styles["CompanyRole"]))
            if dates:
                story.append(Paragraph(dates, styles["SkillText"]))
            bullets = exp.get("bullets", []) or []
            # Show all bullets — the AI prompt already controls count
            # (6-8 for latest experience, 3-5 for others).
            # Prioritize numeric-impact bullets first, then others.
            numeric_bullets = [b for b in bullets if re.search(r"\d", str(b or ""))]
            other_bullets = [b for b in bullets if b not in numeric_bullets]
            chosen = numeric_bullets + other_bullets
            if not chosen:
                chosen = bullets
            for bullet in chosen:
                story.append(Paragraph(f"• {bullet}", styles["ResumeBullet"]))
            story.append(Spacer(1, 4))

    # --- Skills ---
    skills = tailored_resume.get("skills", [])
    if skills:
        story.append(Paragraph("SKILLS", styles["SectionTitle"]))
        story.append(Paragraph(", ".join(skills), styles["SkillText"]))
        story.append(Spacer(1, 6))

    # --- Education ---
    education = tailored_resume.get("education", [])
    if education:
        story.append(Paragraph("EDUCATION", styles["SectionTitle"]))
        for edu in education:
            line = f"{edu.get('degree', '')} — {edu.get('school', '')}"
            year = edu.get("year", "")
            if year:
                line += f" ({year})"
            story.append(Paragraph(line, styles["SkillText"]))

    doc.build(story)
    return output_path


def generate_docx(tailored_resume: Dict, job_title: str = "", company: str = "") -> Path:
    """
    Create a Word DOCX resume from the tailored resume JSON.
    Returns the Path to the generated file.
    """
    RESUMES_TAILORED_DIR.mkdir(parents=True, exist_ok=True)
    filename = f"{uuid.uuid4().hex[:12]}_{_safe_stem(company)}.docx"
    output_path = RESUMES_TAILORED_DIR / filename

    logger.info("DOCX generation start: company=%s title=%s", company, job_title)

    doc = Document()
    # Explicitly set Normal font size using Word units (Pt). Using plain ints can render oddly.
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    summary = tailored_resume.get("summary", "")
    if summary:
        _add_paragraph(doc, summary, size_pt=11)

    experience = tailored_resume.get("experience", [])
    if experience:
        _add_paragraph(doc, "EXPERIENCE", bold=True, size_pt=14)
        for exp in experience:
            role = _clean_text(exp.get("role", ""))
            company_name = _clean_text(exp.get("company", ""))
            dates = _clean_text(exp.get("dates", ""))
            role_line = " - ".join([p for p in [role, company_name] if p])
            _add_paragraph(doc, role_line, bold=True, size_pt=12)
            if dates:
                _add_paragraph(doc, dates, size_pt=10)
            bullets = exp.get("bullets", []) or []
            # Show all bullets — the AI prompt controls count
            # (6-8 for latest experience, 3-5 for others).
            numeric_bullets = [b for b in bullets if re.search(r"\d", str(b or ""))]
            other_bullets = [b for b in bullets if b not in numeric_bullets]
            chosen = numeric_bullets + other_bullets
            if not chosen:
                chosen = bullets
            for bullet in chosen:
                btxt = _clean_text(str(bullet))
                if btxt:
                    # Avoid special unicode bullet characters (can render as � in some Word setups).
                    # Use an ASCII dash for stable display everywhere.
                    _add_paragraph(doc, f"- {btxt}", size_pt=11)

    skills = tailored_resume.get("skills", [])
    if skills:
        _add_paragraph(doc, "SKILLS", bold=True, size_pt=14)
        _add_paragraph(doc, ", ".join(_clean_text(str(s)) for s in skills if _clean_text(str(s))), size_pt=11)

    # ATS systems rely on exact keyword matching. Include explicit ATS keywords
    # returned by the AI tailoring step (if present).
    ats_keywords = tailored_resume.get("ats_keywords_used", [])
    if isinstance(ats_keywords, list) and ats_keywords:
        unique_ats = []
        seen = set()
        for k in ats_keywords:
            kk = _clean_text(str(k))
            if kk and kk not in seen:
                unique_ats.append(kk)
                seen.add(kk)
        if unique_ats:
            _add_paragraph(doc, "ATS KEYWORDS", bold=True, size_pt=14)
            _add_paragraph(doc, ", ".join(unique_ats), size_pt=11)

    education = tailored_resume.get("education", [])
    if education:
        _add_paragraph(doc, "EDUCATION", bold=True, size_pt=14)
        for edu in education:
            degree = _clean_text(edu.get("degree", ""))
            school = _clean_text(edu.get("school", ""))
            year = _clean_text(edu.get("year", ""))
            line = " - ".join([p for p in [degree, school] if p])
            if year:
                line = f"{line} ({year})"
            _add_paragraph(doc, line, size_pt=11)

    doc.save(str(output_path))
    logger.info(
        "DOCX generation complete: path=%s summary=%s experience=%s skills=%s education=%s",
        output_path.name,
        bool(_clean_text(tailored_resume.get("summary", ""))),
        len(experience) if isinstance(experience, list) else 0,
        len(skills) if isinstance(skills, list) else 0,
        len(education) if isinstance(education, list) else 0,
    )
    return output_path
